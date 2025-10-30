"""
Сервис для создания документов Word и PDF.
Форматирует и сохраняет сгенерированный текст в различные форматы.
Поддерживает парсинг Markdown разметки.
"""

import logging
import os
import re
from datetime import datetime
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

logger = logging.getLogger(__name__)


class DocumentService:
    """Класс для создания документов в различных форматах"""
    
    def __init__(self, output_dir: str = 'generated_docs'):
        """
        Инициализация сервиса генерации документов
        
        Args:
            output_dir: Директория для сохранения документов
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Document сервис инициализирован, директория: {output_dir}")
    
    def _parse_markdown_line(self, line: str) -> Tuple[str, str, int]:
        """
        Парсинг markdown строки для определения типа и уровня
        
        Args:
            line: Строка текста с возможной markdown разметкой
        
        Returns:
            Tuple[type, text, level] где:
            - type: тип элемента ('heading', 'list', 'numbered_list', 'text')
            - text: очищенный текст
            - level: уровень заголовка или списка
        """
        line = line.strip()
        
        # Проверка на заголовок
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            return ('heading', text, level)
        
        # Проверка на маркированный список
        list_match = re.match(r'^[\*\-\•]\s+(.+)$', line)
        if list_match:
            text = list_match.group(1).strip()
            return ('list', text, 0)
        
        # Проверка на нумерованный список
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if numbered_match:
            text = numbered_match.group(2).strip()
            return ('numbered_list', text, 0)
        
        return ('text', line, 0)
    
    def _apply_markdown_formatting(self, paragraph, text: str):
        """
        Применение markdown форматирования к тексту параграфа
        
        Args:
            paragraph: Параграф документа Word
            text: Текст с markdown разметкой
        """
        # Сложный паттерн для обнаружения всех типов форматирования
        # Обрабатываем по порядку: жирный, курсив, код
        parts = []
        current_pos = 0
        
        # Паттерн для жирного текста: **text** или __text__
        bold_pattern = r'\*\*(.+?)\*\*|__(.+?)__'
        # Паттерн для курсива: *text* или _text_ (но не внутри слова)
        italic_pattern = r'(?<!\w)\*(?!\*)(.+?)(?<!\*)\*(?!\w)|(?<!\w)_(?!_)(.+?)(?<!_)_(?!\w)'
        # Паттерн для кода: `text`
        code_pattern = r'`(.+?)`'
        
        # Сначала заменяем код на плейсхолдеры, чтобы не конфликтовало с другим форматированием
        code_replacements = []
        for match in re.finditer(code_pattern, text):
            placeholder = f"__CODE_{len(code_replacements)}__"
            code_replacements.append(match.group(1))
            text = text[:match.start()] + placeholder + text[match.end():]
        
        # Теперь обрабатываем весь текст
        pos = 0
        while pos < len(text):
            # Проверяем на жирный текст
            bold_match = re.match(r'\*\*(.+?)\*\*|__(.+?)__', text[pos:])
            if bold_match:
                inner_text = bold_match.group(1) or bold_match.group(2)
                run = paragraph.add_run(inner_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                pos += bold_match.end()
                continue
            
            # Проверяем на курсив
            italic_match = re.match(r'(?<!\w)\*(?!\*)(.+?)(?<!\*)\*(?!\w)|(?<!\w)_(?!_)(.+?)(?<!_)_(?!\w)', text[pos:])
            if italic_match:
                inner_text = italic_match.group(1) or italic_match.group(2)
                run = paragraph.add_run(inner_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
                pos += italic_match.end()
                continue
            
            # Проверяем на плейсхолдер кода
            code_match = re.match(r'__CODE_(\d+)__', text[pos:])
            if code_match:
                code_index = int(code_match.group(1))
                inner_text = code_replacements[code_index]
                run = paragraph.add_run(inner_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
                pos += code_match.end()
                continue
            
            # Обычный текст - находим до следующего форматирования
            next_format = re.search(r'\*\*|__|\*(?!\*)|_(?!_)|__CODE_\d+__', text[pos:])
            if next_format:
                plain_text = text[pos:pos + next_format.start()]
                if plain_text:
                    run = paragraph.add_run(plain_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                pos += next_format.start()
            else:
                # Остаток строки - обычный текст
                plain_text = text[pos:]
                if plain_text:
                    run = paragraph.add_run(plain_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                break
    
    def _process_content_lines(self, doc: Document, content: str):
        """
        Обработка контента с учетом markdown разметки
        
        Args:
            doc: Документ Word
            content: Текстовое содержимое с markdown
        """
        lines = content.split('\n')
        i = 0
        list_items = []
        current_list_type = None
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Пустая строка - завершаем список если он был
                if list_items:
                    self._add_list_to_doc(doc, list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                i += 1
                continue
            
            line_type, text, level = self._parse_markdown_line(line)
            
            if line_type == 'heading':
                # Завершаем список если он был
                if list_items:
                    self._add_list_to_doc(doc, list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                
                # Добавляем заголовок
                heading = doc.add_heading(text, level=min(level + 1, 9))
                if level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif line_type in ('list', 'numbered_list'):
                # Если тип списка изменился, завершаем предыдущий
                if current_list_type and current_list_type != line_type:
                    self._add_list_to_doc(doc, list_items, current_list_type)
                    list_items = []
                
                current_list_type = line_type
                list_items.append(text)
                
            else:
                # Обычный текст
                # Завершаем список если он был
                if list_items:
                    self._add_list_to_doc(doc, list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                
                # Добавляем параграф с форматированием
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._apply_markdown_formatting(para, text)
            
            i += 1
        
        # Завершаем список в конце если он остался
        if list_items:
            self._add_list_to_doc(doc, list_items, current_list_type)
    
    def _add_list_to_doc(self, doc: Document, items: List[str], list_type: str):
        """
        Добавление списка в документ
        
        Args:
            doc: Документ Word
            items: Список элементов
            list_type: Тип списка ('list' или 'numbered_list')
        """
        for i, item in enumerate(items):
            if list_type == 'numbered_list':
                para = doc.add_paragraph(style='List Number')
            else:
                para = doc.add_paragraph(style='List Bullet')
            
            # Очищаем параграф и добавляем с форматированием
            para.clear()
            if list_type == 'numbered_list':
                # Для нумерованного списка добавляем номер
                run = para.add_run(f"{i + 1}. ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                # Для маркированного - маркер
                run = para.add_run("• ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            self._apply_markdown_formatting(para, item)
    
    async def create_word_document(
        self,
        content: str,
        title: str = "Документ",
        user_id: int = 0
    ) -> Optional[str]:
        """
        Создание документа Word (.docx) с поддержкой Markdown
        
        Args:
            content: Текстовое содержимое документа с Markdown разметкой
            title: Заголовок документа
            user_id: ID пользователя для имени файла
        
        Returns:
            Путь к созданному файлу или None в случае ошибки
        """
        try:
            # Создание нового документа
            doc = Document()
            
            # Настройка параметров документа
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1)
            
            # Добавление основного заголовка
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавление даты
            date_paragraph = doc.add_paragraph()
            date_run = date_paragraph.add_run(
                f"Дата создания: {datetime.now().strftime('%d.%m.%Y')}"
            )
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(10)
            date_run.italic = True
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Добавляем разделитель
            doc.add_paragraph()
            
            # Обработка контента с markdown
            self._process_content_lines(doc, content)
            
            # Генерация имени файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"document_{user_id}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Сохранение документа
            doc.save(filepath)
            logger.info(f"Word документ создан: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Ошибка при создании Word документа: {e}")
            return None
    
    async def create_pdf_document(
        self,
        content: str,
        title: str = "Документ",
        user_id: int = 0
    ) -> Optional[str]:
        """
        Создание PDF документа
        
        Args:
            content: Текстовое содержимое документа
            title: Заголовок документа
            user_id: ID пользователя для имени файла
        
        Returns:
            Путь к созданному файлу или None в случае ошибки
        """
        try:
            # Генерация имени файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"document_{user_id}_{timestamp}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            # Создание PDF документа
            doc = SimpleDocTemplate(filepath, pagesize=A4,
                                   rightMargin=inch, leftMargin=inch,
                                   topMargin=inch, bottomMargin=inch)
            
            # Контейнер для элементов документа
            story = []
            
            # Стили
            styles = getSampleStyleSheet()
            
            # Стиль для заголовка
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='black',
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Times-Bold'
            )
            
            # Стиль для основного текста
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=12,
                leading=16,
                alignment=TA_JUSTIFY,
                fontName='Times-Roman',
                spaceAfter=12
            )
            
            # Стиль для подзаголовков
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor='black',
                spaceAfter=12,
                spaceBefore=12,
                fontName='Times-Bold'
            )
            
            # Стиль для списков
            list_style = ParagraphStyle(
                'CustomList',
                parent=styles['BodyText'],
                fontSize=12,
                leading=14,
                leftIndent=20,
                fontName='Times-Roman',
                spaceAfter=6
            )
            
            # Добавление заголовка
            story.append(Paragraph(title, title_style))
            
            # Добавление даты
            date_text = f"<i>Дата создания: {datetime.now().strftime('%d.%m.%Y')}</i>"
            story.append(Paragraph(date_text, body_style))
            story.append(Spacer(1, 0.3 * inch))
            
            # Обработка контента построчно
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.1 * inch))
                    continue
                
                line_type, text, level = self._parse_markdown_line(line)
                
                # Экранирование HTML символов
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Применение markdown форматирования для PDF
                text = self._convert_markdown_to_html(text)
                
                if line_type == 'heading':
                    story.append(Paragraph(text, heading_style))
                elif line_type in ('list', 'numbered_list'):
                    if line_type == 'numbered_list':
                        text = f"• {text}"
                    else:
                        text = f"• {text}"
                    story.append(Paragraph(text, list_style))
                else:
                    story.append(Paragraph(text, body_style))
            
            # Сборка документа
            doc.build(story)
            logger.info(f"PDF документ создан: {filepath}")
            
            return filepath
            
        except Exception as e:
            logger.error(f"Ошибка при создании PDF документа: {e}")
            return None
    
    def _convert_markdown_to_html(self, text: str) -> str:
        """
        Конвертация markdown в HTML для ReportLab
        
        Args:
            text: Текст с markdown разметкой
        
        Returns:
            Текст с HTML тегами
        """
        # Жирный текст
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)
        
        # Курсив
        text = re.sub(r'(?<!\w)\*(?!\*)(.+?)(?<!\*)\*(?!\w)', r'<i>\1</i>', text)
        text = re.sub(r'(?<!\w)_(?!_)(.+?)(?<!_)_(?!\w)', r'<i>\1</i>', text)
        
        # Код (моноширинный)
        text = re.sub(r'`(.+?)`', r'<font name="Courier">\1</font>', text)
        
        return text
    
    def cleanup_file(self, filepath: str) -> bool:
        """
        Удаление временного файла
        
        Args:
            filepath: Путь к файлу для удаления
        
        Returns:
            True если файл удалён успешно, False в противном случае
        """
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                logger.info(f"Файл удалён: {filepath}")
                return True
            return False
        except Exception as e:
            logger.error(f"Ошибка при удалении файла {filepath}: {e}")
            return False
